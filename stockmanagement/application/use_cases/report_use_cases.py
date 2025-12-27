"""Report use cases for generating and downloading reports."""

from __future__ import annotations

import copy
import logging
from datetime import datetime, timedelta
from decimal import Decimal
from typing import Any
from uuid import UUID, uuid4

from django.db import transaction
from django.template.loader import render_to_string
from django.utils import timezone

from application.dto.report_dto import ReportCreateDTO, ReportResponseDTO
from application.use_cases.inventory_use_cases import (
    GenerateInventoryReportUseCase,
    GenerateStockReportUseCase,
)
from application.use_cases.sales_use_cases import GenerateSalesReportUseCase
from domain.business.services import BusinessDomainService
from domain.reports.entities import (
    Report,
    ReportFormat,
    ReportStatus,
    ReportType,
)
from domain.reports.repositories import ReportRepository
from domain.users.repositories import UserRepository
from shared.exceptions.base import BaseAPIException
from shared.exceptions.specific import ForbiddenError, NotFoundError
from shared.services.s3_service import S3Service
from shared.utils.validation import validate_business_access

try:
    from weasyprint import HTML

    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

try:
    from kombu.exceptions import ConnectionError as CeleryConnectionError

    from tasks.report_tasks import generate_report_task
except ImportError:
    generate_report_task = None
    CeleryConnectionError = Exception


def _serialize_report_payload(payload: dict) -> dict:
    """Serialize payload to JSON-friendly types."""

    def _serialize(value):
        if isinstance(value, datetime):
            return value.isoformat()
        if isinstance(value, Decimal):
            return str(value)
        if isinstance(value, list):
            return [_serialize(item) for item in value]
        if isinstance(value, dict):
            return {key: _serialize(val) for key, val in value.items()}
        return value

    return _serialize(copy.deepcopy(payload))


def _sanitize_metadata_for_response(metadata: dict | None) -> dict:
    """Prepare metadata for API responses."""
    if not metadata:
        return {}
    sanitized = dict(metadata)
    sanitized.pop("business_snapshot", None)
    return sanitized


def _business_snapshot_from_entity(business) -> dict:
    """Capture business details for consistent rendering."""
    if not business:
        return {}
    return {
        "name": business.name,
        "address": business.address,
        "phone": business.phone_number,
        "email": business.email,
        "qr_code_url": business.qr_code_url,
        "logo_url": business.logo_url,
    }


def _parse_iso_datetime(value):
    """Convert ISO datetime strings back to datetime objects when possible."""
    if isinstance(value, str):
        try:
            return datetime.fromisoformat(value)
        except ValueError:
            return value
    return value


def _prepare_payload_for_render(payload: dict) -> dict:
    """Prepare payload with datetime objects restored for template rendering."""
    prepared = copy.deepcopy(payload)
    for key in ("period_start", "period_end", "generated_at", "generated_start", "generated_end"):
        if key in prepared:
            prepared[key] = _parse_iso_datetime(prepared[key])
    return prepared


class _Sentinel:
    """Sentinel value to detect if file_url was provided."""


_SENTINEL = _Sentinel()


def _report_to_dto(
    report: Report, file_url: str | None | _Sentinel = _SENTINEL
) -> ReportResponseDTO:
    """
    Convert report entity to DTO (shared utility function).

    Args:
        report: Report entity to convert
        file_url: Optional file URL. If not provided (default), uses report.file_url.
                  If explicitly set to None, uses None.

    Returns:
        ReportResponseDTO instance
    """
    final_file_url = report.file_url if file_url is _SENTINEL else file_url
    return ReportResponseDTO(
        id=report.id,
        business_id=report.business_id,
        report_type=report.report_type.value,
        format=report.format.value,
        status=report.status.value,
        start_date=report.start_date,
        end_date=report.end_date,
        file_url=final_file_url,
        file_size=report.file_size,
        generated_by=report.generated_by,
        error_message=report.error_message,
        metadata=_sanitize_metadata_for_response(report.metadata),
        created_at=report.created_at,
        updated_at=report.updated_at,
    )


def report_to_dto(report: Report, file_url: str | None = None) -> ReportResponseDTO:
    """
    Convert report entity to DTO (public function for external use).

    Args:
        report: Report entity to convert
        file_url: Optional file URL override. If None, uses report.file_url.

    Returns:
        ReportResponseDTO instance
    """
    return _report_to_dto(report, file_url=file_url if file_url is not None else _SENTINEL)


def _render_report_html(
    report_payload: dict, business_info: dict, generated_by_name: str | None
) -> str:
    """Render report HTML using stored payload and business snapshot."""
    payload_for_render = _prepare_payload_for_render(report_payload)
    generated_ts = payload_for_render.get("generated_at") or timezone.now()

    if payload_for_render.get("summary") and not payload_for_render.get("summary_labels"):
        summary = payload_for_render.get("summary", {})
        payload_for_render["summary_labels"] = {
            k: k.replace("_", " ").title() for k in summary.keys()
        }

    context = {
        "report": payload_for_render,
        "business": business_info,
        "generated_by_name": generated_by_name,
        "generated_at": generated_ts,
    }
    return render_to_string("reports/report.html", context)


def _generate_pdf_from_html(html_content: str) -> bytes:
    """Generate PDF bytes from HTML content."""
    if not WEASYPRINT_AVAILABLE:
        raise BaseAPIException(
            detail="WeasyPrint is required for PDF generation",
            code="PDF_GENERATION_UNAVAILABLE",
            status_code=500,
        )
    try:
        return HTML(string=html_content).write_pdf()
    except Exception as exc:
        logger.error("Error generating PDF report: %s", exc, exc_info=True)
        raise BaseAPIException(
            detail=f"Failed to generate PDF report: {str(exc)}",
            code="PDF_GENERATION_ERROR",
            status_code=500,
        ) from exc


def _generate_word_from_payload(
    report_payload: dict, business_info: dict, generated_by_name: str | None
) -> bytes:
    """Generate Word document bytes from report payload."""
    if not DOCX_AVAILABLE:
        raise BaseAPIException(
            detail="python-docx is required for Word document generation",
            code="WORD_GENERATION_UNAVAILABLE",
            status_code=500,
        )
    try:
        from io import BytesIO

        doc = Document()

        # Title
        report_type = report_payload.get("report_type", "Report").title()
        title = doc.add_heading(f"{report_type} Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Business info
        if business_info:
            business_para = doc.add_paragraph()
            business_para.add_run(f"Business: {business_info.get('name', 'N/A')}").bold = True
            if business_info.get("address"):
                doc.add_paragraph(f"Address: {business_info.get('address')}")
            if business_info.get("phone"):
                doc.add_paragraph(f"Phone: {business_info.get('phone')}")

        # Period
        if report_payload.get("period_start") or report_payload.get("period_end"):
            period_para = doc.add_paragraph()
            period_para.add_run("Period: ").bold = True
            period_text = ""
            if report_payload.get("period_start"):
                period_text += f"From {report_payload['period_start']}"
            if report_payload.get("period_end"):
                if period_text:
                    period_text += " "
                period_text += f"to {report_payload['period_end']}"
            period_para.add_run(period_text)

        # Generated by
        if generated_by_name:
            doc.add_paragraph(f"Generated by: {generated_by_name}")

        doc.add_paragraph()  # Spacing

        # Summary statistics
        summary = report_payload.get("summary", {})
        if summary:
            doc.add_heading("Summary", 1)
            summary_labels = report_payload.get("summary_labels", {})
            for key, value in summary.items():
                label = summary_labels.get(key, key.replace("_", " ").title())
                doc.add_paragraph(f"{label}: {value}")

        # Top products (if available)
        if report_payload.get("top_products"):
            doc.add_heading("Top Products", 1)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Product"
            hdr_cells[1].text = "Quantity"
            hdr_cells[2].text = "Revenue"
            hdr_cells[3].text = "Sales"

            for product in report_payload.get("top_products", [])[:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(product.get("product_name", "N/A"))
                row_cells[1].text = str(product.get("total_quantity_sold", 0))
                row_cells[2].text = str(product.get("total_revenue", 0))
                row_cells[3].text = str(product.get("number_of_sales", 0))

        # Top customers (if available)
        if report_payload.get("top_customers"):
            doc.add_heading("Top Customers", 1)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Customer"
            hdr_cells[1].text = "Total Purchases"
            hdr_cells[2].text = "Number of Invoices"

            for customer in report_payload.get("top_customers", [])[:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(customer.get("customer_name", "N/A"))
                row_cells[1].text = str(customer.get("total_purchases", 0))
                row_cells[2].text = str(customer.get("number_of_invoices", 0))

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    except Exception as exc:
        logger.error("Error generating Word document: %s", exc, exc_info=True)
        raise BaseAPIException(
            detail=f"Failed to generate Word document: {str(exc)}",
            code="WORD_GENERATION_ERROR",
            status_code=500,
        ) from exc


class GenerateAndSaveReportUseCase:
    """Use case for generating and saving a report."""

    def __init__(
        self,
        report_repository: ReportRepository,
        invoice_repository,
        invoice_line_repository,
        product_repository,
        stock_movement_repository,
        business_domain_service: BusinessDomainService,
        report_dto: ReportCreateDTO,
        user_id: UUID,
        user_repository: UserRepository | None = None,
    ) -> None:
        """Initialize use case."""
        self.report_repository = report_repository
        self.invoice_repository = invoice_repository
        self.invoice_line_repository = invoice_line_repository
        self.product_repository = product_repository
        self.stock_movement_repository = stock_movement_repository
        self.business_domain_service = business_domain_service
        self.business_id = report_dto.business_id
        self.user_id = user_id
        self.report_type = ReportType(report_dto.report_type)
        self.format = ReportFormat(report_dto.format)
        self.start_date = self._ensure_timezone_aware(report_dto.start_date)
        self.end_date = self._ensure_timezone_aware(report_dto.end_date)
        self.business = None
        self.user_repository = user_repository
        self.generated_by_label: str | None = None

    @staticmethod
    def _ensure_timezone_aware(dt: datetime | None) -> datetime | None:
        """Ensure datetime is timezone-aware."""
        if dt is None:
            return None
        if timezone.is_naive(dt):
            return timezone.make_aware(dt)
        return dt

    def _build_base_payload(self, report_dto, report_type: str, title: str) -> dict:
        """Build base payload structure common to all reports."""
        return {
            "report_type": report_type,
            "title": title,
            "business_id": str(report_dto.business_id),
            "period_start": report_dto.period_start or self.start_date,
            "period_end": report_dto.period_end or self.end_date,
            "generated_at": report_dto.generated_at or timezone.now(),
            "generated_by": self.generated_by_label,
            "generated_start": self.start_date,
            "generated_end": self.end_date,
        }

    @staticmethod
    def _format_summary_key(key: str) -> str:
        """Convert snake_case key to Title Case label."""
        return key.replace("_", " ").title()

    def _build_sales_payload(self, report_dto) -> dict:
        """Build sales report payload."""
        payload = self._build_base_payload(report_dto, "sales", "Sales Report")
        summary = {
            "total_revenue": str(report_dto.total_revenue),
            "total_invoices": report_dto.total_invoices,
            "items_sold": report_dto.total_items_sold,
            "average_invoice_value": str(report_dto.average_invoice_value),
        }
        payload.update(
            {
                "summary": summary,
                "summary_labels": {k: self._format_summary_key(k) for k in summary.keys()},
                "payment_breakdown": [
                    {
                        "payment_method": item.payment_method,
                        "total_amount": str(item.total_amount),
                        "number_of_invoices": item.number_of_invoices,
                    }
                    for item in report_dto.sales_by_payment_method
                ],
                "status_breakdown": [
                    {
                        "status": item.status,
                        "total_amount": str(item.total_amount),
                        "number_of_invoices": item.number_of_invoices,
                    }
                    for item in report_dto.sales_by_status
                ],
                "top_products": [
                    {
                        "product_id": str(p.product_id),
                        "product_name": p.product_name,
                        "total_quantity_sold": p.total_quantity_sold,
                        "total_revenue": str(p.total_revenue),
                        "number_of_sales": p.number_of_sales,
                    }
                    for p in report_dto.top_products
                ],
                "top_customers": [
                    {
                        "customer_id": str(c.customer_id) if c.customer_id else None,
                        "customer_name": c.customer_name,
                        "total_purchases": str(c.total_purchases),
                        "number_of_invoices": c.number_of_invoices,
                    }
                    for c in report_dto.top_customers
                ],
            }
        )
        return payload

    def _build_inventory_payload(self, report_dto) -> dict:
        """Build inventory report payload."""
        payload = self._build_base_payload(report_dto, "inventory", "Inventory Report")
        summary = {
            "total_products": report_dto.total_products,
            "inventory_value": str(report_dto.total_inventory_value),
            "low_stock_items": report_dto.low_stock_products,
            "expired_products": report_dto.expired_products,
        }
        payload.update(
            {
                "summary": summary,
                "summary_labels": {k: self._format_summary_key(k) for k in summary.keys()},
                "products": [
                    {
                        "product_id": str(p.product_id),
                        "product_name": p.product_name,
                        "current_quantity": p.current_quantity,
                        "unit_price": str(p.unit_price),
                        "total_value": str(p.total_value),
                        "is_low_stock": p.is_low_stock,
                        "is_expired": p.is_expired,
                        "expiry_date": p.expiry_date,
                    }
                    for p in report_dto.products
                ],
                "stock_movements_summary": [
                    {
                        "movement_type": s.movement_type,
                        "total_quantity": s.total_quantity,
                        "number_of_movements": s.number_of_movements,
                        "products_affected": s.products_affected,
                    }
                    for s in report_dto.stock_movements_summary
                ],
            }
        )
        return payload

    @staticmethod
    def _build_product_info_dict(product) -> dict:
        """Build product info dictionary (shared across reports)."""
        return {
            "product_id": str(product.product_id),
            "product_name": product.product_name,
            "current_quantity": product.current_quantity,
            "min_quantity": product.min_quantity,
            "unit_price": str(product.unit_price),
            "total_value": str(product.total_value),
            "is_low_stock": product.is_low_stock,
            "is_expired": product.is_expired,
            "expiry_date": product.expiry_date,
        }

    def _build_stock_payload(self, report_dto) -> dict:
        """Build stock report payload."""
        payload = self._build_base_payload(report_dto, "stock", "Stock Report")
        summary = {
            "current_stock_value": str(report_dto.current_stock_value),
            "movements_in": report_dto.stock_movements_in,
            "movements_out": report_dto.stock_movements_out,
            "net_change": report_dto.net_stock_change,
        }
        payload.update(
            {
                "summary": summary,
                "summary_labels": {k: self._format_summary_key(k) for k in summary.keys()},
                "products_by_stock_level": {
                    "low": [
                        self._build_product_info_dict(p)
                        for p in report_dto.products_by_stock_level.get("low", [])
                    ],
                    "normal": [
                        self._build_product_info_dict(p)
                        for p in report_dto.products_by_stock_level.get("normal", [])
                    ],
                    "high": [
                        self._build_product_info_dict(p)
                        for p in report_dto.products_by_stock_level.get("high", [])
                    ],
                },
                "stock_movements_by_type": [
                    {
                        "movement_type": s.movement_type,
                        "total_quantity": s.total_quantity,
                        "number_of_movements": s.number_of_movements,
                        "products_affected": s.products_affected,
                    }
                    for s in report_dto.stock_movements_by_type
                ],
            }
        )
        return payload

    def _validate_report_permission(self) -> None:
        """Validate user has permission to generate the requested report type."""
        permission_map = {
            ReportType.SALES: self.business_domain_service.can_generate_sales_report,
            ReportType.INVENTORY: self.business_domain_service.can_generate_inventory_report,
            ReportType.STOCK: self.business_domain_service.can_generate_stock_report,
        }

        check_permission = permission_map.get(self.report_type)
        if not check_permission:
            raise ValueError(f"Unknown report type: {self.report_type}")

        if not check_permission(self.business_id, self.user_id):
            report_type_name = self.report_type.value.lower()
            raise ForbiddenError(
                detail=f"You don't have permission to generate {report_type_name} reports for this business",
                code="PERMISSION_DENIED",
            )

    @transaction.atomic
    def execute(self) -> ReportResponseDTO:
        """Execute report generation and saving.

        Generates report data immediately and saves it to metadata,
        then dispatches file generation task to Celery (if available).

        Returns:
            ReportResponseDTO with all insights in metadata.payload
        """
        self._validate_report_permission()
        self._ensure_date_range()
        self.generated_by_label = self._get_generated_by_label()

        business = self.business_domain_service.get_business(self.business_id)
        if not business:
            raise NotFoundError(
                detail="Business not found",
                code="BUSINESS_NOT_FOUND",
            )
        self.business = business

        report_id = uuid4()
        metadata = {}
        if self.generated_by_label:
            metadata["generated_by_name"] = self.generated_by_label

        report = Report(
            id=report_id,
            business_id=self.business_id,
            report_type=self.report_type,
            format=self.format,
            status=ReportStatus.GENERATING,
            start_date=self.start_date,
            end_date=self.end_date,
            file_path=None,
            file_url=None,
            file_size=None,
            generated_by=self.user_id,
            error_message=None,
            metadata=metadata,
            created_at=timezone.now(),
            updated_at=timezone.now(),
        )

        report = self.report_repository.create(report)

        try:
            if self.report_type == ReportType.SALES:
                report_payload = self._generate_sales_report()
            elif self.report_type == ReportType.INVENTORY:
                report_payload = self._generate_inventory_report()
            elif self.report_type == ReportType.STOCK:
                report_payload = self._generate_stock_report()
            else:
                raise ValueError(f"Unknown report type: {self.report_type}")

            serialized_payload = _serialize_report_payload(report_payload)
            business_snapshot = _business_snapshot_from_entity(self.business)

            updated_metadata: dict[str, Any] = report.metadata or {}
            updated_metadata["payload"] = serialized_payload
            updated_metadata["business_snapshot"] = business_snapshot
            if self.generated_by_label:
                updated_metadata["generated_by_name"] = self.generated_by_label

            report.metadata = updated_metadata
            report.updated_at = timezone.now()
            report = self.report_repository.update(report)

            if generate_report_task:
                try:
                    start_date_str = self.start_date.isoformat() if self.start_date else None
                    end_date_str = self.end_date.isoformat() if self.end_date else None

                    generate_report_task.delay(
                        report_id=report_id,
                        business_id=self.business_id,
                        user_id=self.user_id,
                        report_type=self.report_type.value,
                        format=self.format.value,
                        start_date=start_date_str,
                        end_date=end_date_str,
                    )
                    logger.info(
                        f"Report {report_id} file generation task dispatched to Celery "
                        f"(type: {self.report_type.value}, format: {self.format.value})"
                    )
                except (CeleryConnectionError, Exception) as celery_error:
                    logger.warning(
                        f"Celery unavailable, falling back to synchronous file generation for report {report_id}: {celery_error}",
                        exc_info=True,
                    )
                    return self._generate_and_upload_file(
                        report_id, report_payload, business_snapshot
                    )
            else:
                logger.info(
                    f"Generating report {report_id} file synchronously (Celery not available) "
                    f"(type: {self.report_type.value}, format: {self.format.value})"
                )
                return self._generate_and_upload_file(report_id, report_payload, business_snapshot)

            return _report_to_dto(report)

        except Exception as e:
            logger.error(
                f"Failed to initiate report generation for {report_id}: {str(e)}",
                exc_info=True,
                extra={"report_id": str(report_id)},
            )

            report.status = ReportStatus.FAILED
            report.error_message = f"Failed to initiate report generation: {str(e)}"
            report.updated_at = timezone.now()
            self.report_repository.update(report)

            raise

    def _generate_and_upload_file(
        self, report_id: UUID, report_payload: dict, business_snapshot: dict
    ) -> ReportResponseDTO:
        """Generate and upload report file to S3.

        Args:
            report_id: ID of the report
            report_payload: Pre-generated report payload with all insights
            business_snapshot: Business information snapshot

        Returns:
            ReportResponseDTO with completed report information
        """
        report = self.report_repository.get_by_id(report_id)
        if not report:
            raise NotFoundError(
                detail="Report not found",
                code="REPORT_NOT_FOUND",
            )

        if not self.generated_by_label:
            self.generated_by_label = self._get_generated_by_label()

        file_url = None
        file_size = None
        try:
            s3_service = S3Service()
            filename_base = f"{self.report_type.value}-{report_id}"
            file_bytes = None

            if self.format == ReportFormat.PDF:
                html_content = _render_report_html(
                    report_payload=report_payload,
                    business_info=business_snapshot,
                    generated_by_name=self.generated_by_label,
                )
                file_bytes = _generate_pdf_from_html(html_content)
                file_url = s3_service.upload_pdf(file_bytes, filename=filename_base)
            elif self.format == ReportFormat.WORD:
                file_bytes = _generate_word_from_payload(
                    report_payload=report_payload,
                    business_info=business_snapshot,
                    generated_by_name=self.generated_by_label,
                )
                file_url = s3_service.upload_word(file_bytes, filename=filename_base)
            else:
                logger.warning(
                    f"Unsupported format {self.format.value} for report {report_id}",
                    extra={"report_id": str(report_id), "format": self.format.value},
                )

            if file_url and file_bytes:
                file_size = len(file_bytes)
                logger.info(
                    f"File uploaded to S3 for report {report_id}: {file_url}, size: {file_size} bytes",
                    extra={
                        "report_id": str(report_id),
                        "file_url": file_url,
                        "file_size": file_size,
                    },
                )
        except Exception as file_error:
            logger.error(
                f"Failed to generate/upload file for report {report_id}: {file_error}",
                exc_info=True,
                extra={"report_id": str(report_id)},
            )

        report.status = ReportStatus.COMPLETED
        report.file_path = None
        report.file_url = file_url
        report.file_size = file_size
        report.updated_at = timezone.now()
        report = self.report_repository.update(report)

        return _report_to_dto(report)

    def generate_and_save_report_content(self, report_id: UUID) -> ReportResponseDTO:
        """Generate report file and upload to S3 (called by Celery task).

        This method assumes report data (payload) is already generated and saved
        in metadata. It only generates the file and uploads it.

        Args:
            report_id: ID of the report to generate file for

        Returns:
            ReportResponseDTO with completed report information
        """
        report = self.report_repository.get_by_id(report_id)
        if not report:
            raise NotFoundError(
                detail="Report not found",
                code="REPORT_NOT_FOUND",
            )

        if not self.business:
            self.business = self.business_domain_service.get_business(self.business_id)
            if not self.business:
                raise NotFoundError(
                    detail="Business not found",
                    code="BUSINESS_NOT_FOUND",
                )

        if not self.generated_by_label:
            self.generated_by_label = self._get_generated_by_label()

        try:
            # Retrieve payload from metadata (already generated in execute())
            metadata = report.metadata or {}
            payload_data = metadata.get("payload")
            business_snapshot = metadata.get("business_snapshot") or _business_snapshot_from_entity(
                self.business
            )

            if not payload_data:
                logger.warning(
                    f"Payload not found in metadata for report {report_id}, regenerating...",
                    extra={"report_id": str(report_id)},
                )
                if self.report_type == ReportType.SALES:
                    report_payload = self._generate_sales_report()
                elif self.report_type == ReportType.INVENTORY:
                    report_payload = self._generate_inventory_report()
                elif self.report_type == ReportType.STOCK:
                    report_payload = self._generate_stock_report()
                else:
                    raise ValueError(f"Unknown report type: {self.report_type}")
            else:
                # Deserialize payload for file generation (convert ISO strings back to datetime)
                report_payload = _prepare_payload_for_render(payload_data)

            return self._generate_and_upload_file(report_id, report_payload, business_snapshot)

        except Exception as e:
            logger.error(f"Error generating report {report_id}: {str(e)}", exc_info=True)
            report.status = ReportStatus.FAILED
            report.error_message = str(e)
            report.updated_at = timezone.now()
            self.report_repository.update(report)
            raise BaseAPIException(
                detail=f"Failed to generate report: {str(e)}",
                code="REPORT_GENERATION_ERROR",
                status_code=500,
            ) from e

    def _generate_sales_report(self) -> dict:
        """Generate sales report payload."""
        use_case = GenerateSalesReportUseCase(
            invoice_repository=self.invoice_repository,
            invoice_line_repository=self.invoice_line_repository,
            product_repository=self.product_repository,
            business_domain_service=self.business_domain_service,
            business_id=self.business_id,
            user_id=self.user_id,
            start_date=self.start_date,
            end_date=self.end_date,
        )
        report_dto = use_case.execute()
        return self._build_sales_payload(report_dto)

    def _generate_inventory_report(self) -> dict:
        """Generate inventory report payload."""
        use_case = GenerateInventoryReportUseCase(
            product_repository=self.product_repository,
            stock_movement_repository=self.stock_movement_repository,
            business_domain_service=self.business_domain_service,
            business_id=self.business_id,
            user_id=self.user_id,
            start_date=self.start_date,
            end_date=self.end_date,
        )
        report_dto = use_case.execute()
        return self._build_inventory_payload(report_dto)

    def _generate_stock_report(self) -> dict:
        """Generate stock report payload."""
        use_case = GenerateStockReportUseCase(
            stock_movement_repository=self.stock_movement_repository,
            product_repository=self.product_repository,
            business_domain_service=self.business_domain_service,
            business_id=self.business_id,
            user_id=self.user_id,
            start_date=self.start_date,
            end_date=self.end_date,
        )
        report_dto = use_case.execute()
        return self._build_stock_payload(report_dto)

    def _ensure_date_range(self) -> None:
        """Ensure start and end dates have sensible defaults."""
        now = timezone.now()
        default_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        default_end = default_start + timedelta(days=1) - timedelta(microseconds=1)

        if self.start_date is None:
            self.start_date = default_start
        if self.end_date is None:
            self.end_date = default_end

        if self.start_date > self.end_date:
            self.start_date, self.end_date = self.end_date, self.start_date

    def _get_generated_by_label(self) -> str:
        """Resolve the name/email of the user who generated the report."""
        if not self.user_repository:
            return str(self.user_id)
        user = self.user_repository.get_by_id(self.user_id)
        if not user:
            return str(self.user_id)
        return user.name or user.email or str(user.id)


class DownloadReportUseCase:
    """Use case for downloading a report file from S3."""

    def __init__(
        self,
        report_repository: ReportRepository,
        business_domain_service: BusinessDomainService,
        report_id: UUID,
        user_id: UUID,
    ) -> None:
        """Initialize use case."""
        self.report_repository = report_repository
        self.business_domain_service = business_domain_service
        self.report_id = report_id
        self.user_id = user_id

    def execute(self) -> tuple[bytes, str, str]:
        """Download report file from S3.

        Returns:
            Tuple of (content_bytes, content_type, filename)
        """
        report = self.report_repository.get_by_id(self.report_id)
        if not report:
            raise NotFoundError(
                detail="Report not found",
                code="REPORT_NOT_FOUND",
            )

        if report.status != ReportStatus.COMPLETED:
            raise BaseAPIException(
                detail=f"Report is not ready. Status: {report.status.value}",
                code="REPORT_NOT_READY",
                status_code=400,
            )

        validate_business_access(
            business_domain_service=self.business_domain_service,
            business_id=report.business_id,
            user_id=self.user_id,
            error_message="You don't have access to this report",
        )

        if not report.file_url:
            raise BaseAPIException(
                detail="Report file not available. The file may still be generating.",
                code="FILE_NOT_AVAILABLE",
                status_code=404,
            )

        report_format = report.format.value

        # Download file from S3
        try:
            s3_service = S3Service()
            key = s3_service._extract_key_from_s3_url(report.file_url)
            if not key:
                raise BaseAPIException(
                    detail="Invalid S3 URL format",
                    code="INVALID_S3_URL",
                    status_code=500,
                )

            s3_client = s3_service.s3_client
            response = s3_client.get_object(Bucket=s3_service.aws_bucket_name, Key=key)
            file_bytes = response["Body"].read()
        except Exception as e:
            logger.error(
                f"Failed to download file from S3 for report {self.report_id}: {e}",
                exc_info=True,
                extra={"report_id": str(self.report_id), "file_url": report.file_url},
            )
            raise BaseAPIException(
                detail="Failed to download report file. Please try again later.",
                code="DOWNLOAD_FAILED",
                status_code=500,
            ) from e

        if report_format == "pdf":
            content_type = "application/pdf"
            extension = "pdf"
        elif report_format == "word":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            extension = "docx"
        else:
            content_type = "application/octet-stream"
            extension = report_format

        created_at = getattr(report, "created_at", None)
        timestamp_str = created_at.strftime("%Y%m%d_%H%M%S") if created_at else "report"
        filename = f"{report.report_type.value}_report_{timestamp_str}.{extension}"

        return file_bytes, content_type, filename


class DeleteReportUseCase:
    """Use case for deleting a report."""

    def __init__(
        self,
        report_repository: ReportRepository,
        business_domain_service: BusinessDomainService,
        report_id: UUID,
        user_id: UUID,
    ) -> None:
        """Initialize use case."""
        self.report_repository = report_repository
        self.business_domain_service = business_domain_service
        self.report_id = report_id
        self.user_id = user_id

    def execute(self) -> None:
        """Delete the report and its file from S3."""
        report = self.report_repository.get_by_id(self.report_id)
        if not report:
            raise NotFoundError(
                detail="Report not found",
                code="REPORT_NOT_FOUND",
            )

        validate_business_access(
            business_domain_service=self.business_domain_service,
            business_id=report.business_id,
            user_id=self.user_id,
            error_message="You don't have access to this report",
        )

        if report.file_url:
            try:
                s3_service = S3Service()
                s3_service.delete_file_safe(report.file_url)
            except Exception as e:
                logger.warning(
                    f"Failed to delete file from S3 for report {self.report_id}: {e}",
                    extra={"report_id": str(self.report_id), "file_url": report.file_url},
                )

        self.report_repository.delete(report.id)


class ListReportsUseCase:
    """Use case for listing reports for a business."""

    def __init__(
        self,
        report_repository: ReportRepository,
        business_domain_service: BusinessDomainService,
        business_id: UUID,
        user_id: UUID,
        report_type: str | None = None,
        status: str | None = None,
        start_date: datetime | None = None,
        end_date: datetime | None = None,
        limit: int = 100,
    ) -> None:
        """Initialize use case."""
        self.report_repository = report_repository
        self.business_domain_service = business_domain_service
        self.business_id = business_id
        self.user_id = user_id
        self.report_type = ReportType(report_type) if report_type else None
        self.status = ReportStatus(status) if status else None
        self.start_date = start_date
        self.end_date = end_date
        self.limit = limit

    def execute(self) -> list[ReportResponseDTO]:
        """Return reports after verifying access."""
        validate_business_access(
            business_domain_service=self.business_domain_service,
            business_id=self.business_id,
            user_id=self.user_id,
        )

        reports = self.report_repository.get_by_business(
            business_id=self.business_id,
            report_type=self.report_type.value if self.report_type else None,
            status=self.status.value if self.status else None,
            start_date=self.start_date,
            end_date=self.end_date,
            limit=self.limit,
        )
        return [_report_to_dto(report) for report in reports]
